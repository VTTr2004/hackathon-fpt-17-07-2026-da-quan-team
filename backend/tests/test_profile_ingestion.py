from pathlib import Path

from docx import Document as DocxDocument

from app.modules.document_chatbot.ingestion import text_to_chunks
from app.modules.profile_ingestion.extractor import select_relevant_blocks, validate_llm_result
from app.modules.profile_ingestion.normalizers import normalize_profile_value
from app.modules.profile_ingestion.service import _document_blocks
from app.modules.profile_ingestion.schemas import (
    EvidenceBlock,
    LLMCandidate,
    LLMEvidenceReference,
    LLMExtractionResult,
)
from app.services.document_parser import extract_text, has_extractable_text
from app.services.ocr_service import normalize_ocr_text


def test_profile_value_normalization_is_deterministic() -> None:
    assert normalize_profile_value("stage", "series-a") == "Series A"
    assert normalize_profile_value("target_customers", "Sinh viên, Freelancer\nSinh viên") == [
        "Sinh viên",
        "Freelancer",
    ]


def test_validated_candidate_requires_quote_from_real_block() -> None:
    block = EvidenceBlock(
        block_id="doc-1:chunk:0",
        document_id="doc-1",
        filename="pitch.pdf",
        text="Startup hiện đang ở giai đoạn Seed và bắt đầu mở rộng thị trường.",
        metadata={"page": 3},
    )
    result = LLMExtractionResult(
        candidates=[
            LLMCandidate(
                field_key="stage",
                proposed_value="Seed",
                extraction_status="found",
                evidence=[
                    LLMEvidenceReference(
                        block_id=block.block_id,
                        quote="hiện đang ở giai đoạn Seed",
                    )
                ],
            )
        ]
    )

    candidate = validate_llm_result(result, [block], ["stage"])[0]

    assert candidate.status == "found"
    assert candidate.proposed_value == "Seed"
    assert candidate.evidence[0]["page"] == 3
    assert candidate.confidence >= 0.8


def test_hallucinated_quote_cannot_become_found_candidate() -> None:
    block = EvidenceBlock(
        block_id="doc-1:chunk:0",
        document_id="doc-1",
        filename="pitch.pdf",
        text="Tài liệu chỉ mô tả sản phẩm.",
        metadata={"page": 1},
    )
    result = LLMExtractionResult(
        candidates=[
            LLMCandidate(
                field_key="traction",
                proposed_value="10.000 khách hàng",
                extraction_status="found",
                evidence=[
                    LLMEvidenceReference(block_id=block.block_id, quote="Đã có 10.000 khách hàng")
                ],
            )
        ]
    )

    candidate = validate_llm_result(result, [block], ["traction"])[0]

    assert candidate.status == "ambiguous"
    assert candidate.evidence == []
    assert any("Quote không khớp" in warning for warning in candidate.warnings)


def test_retrieval_prefers_blocks_with_field_keywords() -> None:
    blocks = [
        EvidenceBlock(
            block_id="generic",
            document_id="doc-1",
            filename="pitch.pdf",
            text="Thông tin chung về doanh nghiệp.",
        ),
        EvidenceBlock(
            block_id="problem",
            document_id="doc-1",
            filename="pitch.pdf",
            text="Vấn đề của khách hàng là thiếu không gian học tập yên tĩnh.",
        ),
    ]

    selected = select_relevant_blocks(blocks, ["problem"], per_field=1, maximum=1)

    assert [block.block_id for block in selected] == ["problem"]


def test_text_chunks_do_not_cross_source_markers() -> None:
    chunks = text_to_chunks(
        "[PAGE 1]\nNội dung trang một.\n[PAGE 2]\nNội dung trang hai.",
        document_id="doc-1",
        filename="pitch.pdf",
    )

    assert [chunk["metadata"]["page"] for chunk in chunks] == [1, 2]
    assert "trang hai" not in chunks[0]["text"]


def test_source_markers_without_content_are_not_extractable() -> None:
    assert not has_extractable_text("[PAGE 1]\n\n[PAGE 2]\n")
    assert has_extractable_text("[PAGE 1]\nNội dung có thể chọn và sao chép")


def test_ocr_output_normalizes_page_headings_and_code_fences() -> None:
    assert normalize_ocr_text("```markdown\n# Page 1\nXin chào\n```") == "[PAGE 1]\nXin chào"
    assert normalize_ocr_text("Nội dung ảnh") == "[PAGE 1]\nNội dung ảnh"


def test_docx_tables_become_locatable_evidence_blocks(tmp_path: Path) -> None:
    path = tmp_path / "profile.docx"
    document = DocxDocument()
    document.add_paragraph("Giới thiệu doanh nghiệp")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Chỉ tiêu"
    table.cell(0, 1).text = "Giá trị"
    table.cell(1, 0).text = "Giai đoạn"
    table.cell(1, 1).text = "Seed"
    document.save(path)

    chunks = text_to_chunks(extract_text(path), document_id="doc-1", filename=path.name)

    table_chunk = next(chunk for chunk in chunks if chunk["metadata"].get("table") == 1)
    assert "Giai đoạn\tSeed" in table_chunk["text"]


def test_document_chunks_map_chunk_id_to_evidence_block_id(tmp_path: Path) -> None:
    path = tmp_path / "profile.json"
    path.write_text('{"name": "Góc Hồ Coffee"}', encoding="utf-8")

    blocks = _document_blocks(
        {"id": "doc-1", "filename": path.name, "storage_path": str(path), "text": ""}
    )

    assert blocks
    assert blocks[0].block_id
    assert blocks[0].document_id == "doc-1"
